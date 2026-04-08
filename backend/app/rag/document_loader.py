"""
Document loader and text splitter.

Uses the DocumentLoaderRegistry for pluggable document format support.
"""

from typing import List, Optional
from pathlib import Path
from langchain_core.documents import Document
from langchain_text_splitters import RecursiveCharacterTextSplitter, MarkdownHeaderTextSplitter

from app.core.logging_config import get_logger
from app.rag.loaders.registry import DocumentLoaderRegistry
from app.rag.utils import is_valid_text, detect_encoding

logger = get_logger("document_loader")


class DocumentLoader:
    """
    Document loader that delegates to registered loaders.

    This class maintains backward compatibility and delegates
    to the DocumentLoaderRegistry for actual loading.
    """

    def __init__(self, encoding: str = "utf-8", fast_mode: bool = True):
        self.encoding = encoding
        self.fast_mode = fast_mode

    def load_file(self, file_path: str) -> List[Document]:
        """
        Load a document from file.

        Delegates to the appropriate loader registered in DocumentLoaderRegistry.

        Args:
            file_path: Path to the file to load

        Returns:
            List of LangChain Document objects
        """
        path = Path(file_path)
        suffix = path.suffix.lower()

        # Check if we have a loader for this extension
        try:
            loader = DocumentLoaderRegistry.get_loader(file_path)
            docs = loader.load(file_path)

            if docs:
                return docs
        except ValueError as e:
            logger.warning(f"No loader for {suffix}: {e}")
        except Exception as e:
            logger.error(f"Error loading file {file_path}: {e}")
            return []

        # Fall back to built-in loaders for backward compatibility
        if suffix == ".pdf":
            return self._load_pdf_fallback(file_path)
        elif suffix in [".md", ".markdown"]:
            return self._load_markdown_fallback(file_path)
        elif suffix == ".txt":
            return self._load_text_fallback(file_path)
        elif suffix in [".doc", ".docx"]:
            return self._load_docx_fallback(file_path)
        else:
            logger.warning(f"Unsupported file type: {suffix}")
            return []

    def _load_pdf_fallback(self, file_path: str) -> List[Document]:
        """Fallback PDF loader using PyMuPDF."""
        from app.rag.loaders.pdf_loader import PDFLoader
        loader = PDFLoader()
        return loader.load(file_path)

    def _load_markdown_fallback(self, file_path: str) -> List[Document]:
        """Fallback markdown loader."""
        from langchain_community.document_loaders import TextLoader

        encodings_to_try = ['utf-8', 'gbk', 'gb2312', 'gb18030']

        for encoding in encodings_to_try:
            try:
                loader = TextLoader(file_path, encoding=encoding)
                docs = loader.load()

                if docs and docs[0].page_content:
                    content = docs[0].page_content
                    if is_valid_text(content):
                        path = Path(file_path)
                        for doc in docs:
                            doc.metadata["file_path"] = str(file_path)
                            doc.metadata["file_name"] = path.name
                            doc.metadata["file_type"] = "md"
                        return docs
            except UnicodeDecodeError:
                continue
            except Exception as e:
                logger.debug(f"Error with encoding {encoding}: {e}")
                continue

        return []

    def _load_text_fallback(self, file_path: str) -> List[Document]:
        """Fallback text loader."""
        from langchain_community.document_loaders import TextLoader

        encodings = ['utf-8', 'gbk', 'gb2312', 'gb18030', 'big5', 'utf-16']
        detected = detect_encoding(file_path)
        if detected and detected.lower() not in [e.lower() for e in encodings]:
            encodings.insert(0, detected)

        for encoding in encodings:
            try:
                loader = TextLoader(file_path, encoding=encoding)
                docs = loader.load()

                if docs and docs[0].page_content:
                    content = docs[0].page_content
                    if is_valid_text(content):
                        path = Path(file_path)
                        for doc in docs:
                            doc.metadata["file_path"] = str(file_path)
                            doc.metadata["file_name"] = path.name
                            doc.metadata["file_type"] = "txt"
                        return docs
            except UnicodeDecodeError:
                continue
            except Exception as e:
                logger.debug(f"Error with encoding {encoding}: {e}")
                continue

        return []

    def _load_docx_fallback(self, file_path: str) -> List[Document]:
        """Fallback DOCX loader."""
        try:
            from docx import Document as DocxDocument

            doc = DocxDocument(file_path)
            content = "\n".join([para.text for para in doc.paragraphs if para.text])

            if not is_valid_text(content):
                return []

            path = Path(file_path)
            return [Document(
                page_content=content,
                metadata={
                    "file_path": str(file_path),
                    "file_name": path.name,
                    "file_type": "docx",
                }
            )]
        except ImportError:
            logger.warning("python-docx not installed")
            return []
        except Exception as e:
            logger.error(f"Failed to load docx: {e}")
            return []

    def load_directory(self, directory: str, glob_pattern: str = "**/*") -> List[Document]:
        """
        Load all supported documents from a directory.

        Args:
            directory: Directory path
            glob_pattern: Glob pattern for file matching

        Returns:
            List of all loaded documents
        """
        path = Path(directory)
        if not path.exists() or not path.is_dir():
            logger.warning(f"Directory not found: {directory}")
            return []

        all_docs = []
        supported_extensions = DocumentLoaderRegistry.list_extensions()

        for ext in supported_extensions:
            # Remove the leading dot for glob pattern
            if ext.startswith('.'):
                ext_pattern = f"*{ext}"
            else:
                ext_pattern = f"*.{ext}"

            for file_path in path.glob(f"{glob_pattern}{ext_pattern}"):
                docs = self.load_file(str(file_path))
                all_docs.extend(docs)

        logger.info(f"Loaded {len(all_docs)} documents from {directory}")
        return all_docs


class TextSplitter:
    """
    Text splitter for document chunking.

    Supports different strategies based on file type.
    """

    def __init__(
        self,
        chunk_size: int = 500,
        chunk_overlap: int = 50,
        splitting_strategies: Optional[dict] = None,
    ):
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap
        self.strategies = splitting_strategies or {}

    def split_documents(self, documents: List[Document]) -> List[Document]:
        """Split documents into chunks."""
        if not documents:
            return []

        all_splits = []

        for doc in documents:
            file_type = doc.metadata.get("file_type", "txt")
            splits = self._split_by_type(doc, file_type)
            all_splits.extend(splits)

        logger.info(f"Split {len(documents)} documents into {len(all_splits)} chunks")
        return all_splits

    def _split_by_type(self, doc: Document, file_type: str) -> List[Document]:
        """Split document based on its type."""
        if file_type in ["md", "markdown"]:
            return self._split_markdown(doc)
        elif file_type == "pdf":
            return self._split_pdf(doc)
        else:
            return self._split_text(doc)

    def _split_markdown(self, doc: Document) -> List[Document]:
        """Split markdown document by headers."""
        headers_to_split_on = [
            ("#", "header1"),
            ("##", "header2"),
            ("###", "header3"),
            ("####", "header4"),
        ]

        markdown_splitter = MarkdownHeaderTextSplitter(
            headers_to_split_on=headers_to_split_on,
            strip_headers=False,
        )

        md_splits = markdown_splitter.split_text(doc.page_content)

        text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=self.chunk_size,
            chunk_overlap=self.chunk_overlap,
            separators=["\n\n", "\n", " ", ""],
        )

        final_splits = []
        for split in md_splits:
            if len(split.page_content) > self.chunk_size:
                sub_splits = text_splitter.split_documents([split])
                final_splits.extend(sub_splits)
            else:
                final_splits.append(split)

        for split in final_splits:
            split.metadata.update(doc.metadata)

        return final_splits

    def _split_pdf(self, doc: Document) -> List[Document]:
        """Split PDF document."""
        text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=self.chunk_size or 800,
            chunk_overlap=self.chunk_overlap or 100,
            separators=["\n\n\n", "\n\n", "\n", " ", ""],
        )

        splits = text_splitter.split_documents([doc])
        for split in splits:
            split.metadata.update(doc.metadata)

        return splits

    def _split_text(self, doc: Document) -> List[Document]:
        """Split plain text document."""
        text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=self.chunk_size,
            chunk_overlap=self.chunk_overlap,
            separators=["\n\n", "\n", " ", ""],
        )

        splits = text_splitter.split_documents([doc])
        for split in splits:
            split.metadata.update(doc.metadata)

        return splits
