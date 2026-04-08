"""
DOCX document loader.
"""

from typing import List
from pathlib import Path
from langchain_core.documents import Document

from app.rag.loaders.registry import register_document_loader
from app.rag.utils import is_valid_text
from app.core.logging_config import get_logger

logger = get_logger("docx_loader")


@register_document_loader
class DocxLoader:
    """Microsoft Word DOCX document loader."""

    @property
    def supported_extensions(self) -> List[str]:
        return [".docx"]

    def load(self, file_path: str) -> List[Document]:
        """Load DOCX file and extract paragraph text."""
        try:
            from docx import Document as DocxDocument

            doc = DocxDocument(file_path)
            content = "\n".join([para.text for para in doc.paragraphs if para.text])

            if not is_valid_text(content):
                logger.warning(f"Invalid content in docx file: {file_path}")
                return []

            path = Path(file_path)
            result = [Document(
                page_content=content,
                metadata={
                    "file_path": str(file_path),
                    "file_name": path.name,
                    "file_type": "docx",
                }
            )]

            logger.info(f"Loaded docx file: {file_path}")
            return result

        except ImportError:
            logger.warning("python-docx not installed, cannot load .docx files")
            return []
        except Exception as e:
            logger.error(f"Failed to load docx file {file_path}: {e}")
            return []
